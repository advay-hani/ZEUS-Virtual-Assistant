"""
Document processing module for Zeus Virtual Assistant.
Handles document upload, text extraction, chunking, and storage with comprehensive error handling.
"""
import os
import logging
import re
import time
from typing import Optional, List, Tuple, Callable, Dict, Any
from datetime import datetime
from tkinter import filedialog, messagebox
import PyPDF2
from docx import Document as DocxDocument

from models.data_models import Document
from core.persistence import PersistenceManager
from core.error_handler import ErrorHandler, ErrorCategory, ErrorSeverity, handle_errors, ProgressContext
from core.performance_monitor import get_performance_monitor
from core.memory_optimizer import get_memory_optimizer
from core.background_processor import get_background_processor


class DocumentProcessor:
    """Handles document upload, text extraction, chunking, and storage with comprehensive error handling."""
    
    SUPPORTED_FORMATS = {'.pdf', '.doc', '.docx', '.txt'}
    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB limit
    DEFAULT_CHUNK_SIZE = 1000  # characters per chunk
    CHUNK_OVERLAP = 200  # overlap between chunks
    
    def __init__(self, persistence_manager: Optional[PersistenceManager] = None, 
                 error_handler: Optional[ErrorHandler] = None):
        """Initialize the document processor."""
        self.current_document: Optional[Document] = None
        self.stored_documents: List[Document] = []
        self.persistence_manager = persistence_manager or PersistenceManager()
        self.error_handler = error_handler
        self.logger = logging.getLogger(__name__)
        
        # Performance monitoring
        self.performance_monitor = get_performance_monitor()
        self.memory_optimizer = get_memory_optimizer()
        self.background_processor = get_background_processor()
        
        # Document processing optimization
        self.max_documents_in_memory = 5
        self.chunk_processing_batch_size = 100
        
        # Load existing documents on initialization
        self._load_stored_documents()
        
        # Register optimization callbacks
        self.memory_optimizer.add_optimization_callback(self._handle_memory_optimization)
    
    def upload_document(self) -> Optional[Document]:
        """
        Open file dialog and upload a document with progress indication and error handling.
        
        Returns:
            Document object if successful, None if cancelled or failed
        """
        file_path = None
        progress = None
        
        try:
            # Open file dialog
            file_path = filedialog.askopenfilename(
                title="Select Document",
                filetypes=[
                    ("All Supported", "*.pdf;*.doc;*.docx"),
                    ("PDF files", "*.pdf"),
                    ("Word documents", "*.doc;*.docx"),
                    ("All files", "*.*")
                ]
            )
            
            if not file_path:
                return None  # User cancelled
            
            # Show progress indicator
            if self.error_handler:
                progress = self.error_handler.create_progress_indicator("Processing Document")
                progress.show("Validating document...", indeterminate=True)
            
            # Validate file
            validation_result = self._validate_file(file_path)
            if not validation_result[0]:
                if self.error_handler:
                    self.error_handler.handle_error(
                        ValueError(validation_result[1]),
                        ErrorCategory.DOCUMENT,
                        ErrorSeverity.ERROR,
                        custom_message=f"Document validation failed: {validation_result[1]}",
                        custom_recovery=[
                            "Choose a different document",
                            "Check if the file is corrupted",
                            "Ensure the file format is supported (PDF, DOC, DOCX)"
                        ]
                    )
                else:
                    messagebox.showerror("Invalid File", validation_result[1])
                return None
            
            filename = os.path.basename(file_path)
            file_size = os.path.getsize(file_path)
            
            # Update progress
            if progress:
                progress.update_message(f"Extracting text from {filename}...")
            
            # Extract text with timeout for large files
            text_content = self.extract_text(file_path)
            if not text_content or not text_content.strip():
                error_msg = "Could not extract text from the document. The file may be empty, corrupted, or password-protected."
                if self.error_handler:
                    self.error_handler.handle_error(
                        ValueError("No text extracted"),
                        ErrorCategory.DOCUMENT,
                        ErrorSeverity.ERROR,
                        custom_message=error_msg,
                        custom_recovery=[
                            "Try a different document",
                            "Check if the document is password-protected",
                            "Ensure the document contains readable text"
                        ]
                    )
                else:
                    messagebox.showerror("Extraction Failed", error_msg)
                return None
            
            # Update progress
            if progress:
                progress.update_message("Creating document chunks...")
            
            # Create document object
            document = Document(
                filename=filename,
                file_path=file_path,
                text_content=text_content,
                file_size=file_size
            )
            
            # Chunk the document
            document.chunks = self.chunk_document(text_content)
            
            # Update progress
            if progress:
                progress.update_message("Saving document...")
            
            # Store the document
            if not self._store_document(document):
                error_msg = "Failed to save the document. Please check disk space and permissions."
                if self.error_handler:
                    self.error_handler.handle_error(
                        IOError("Document storage failed"),
                        ErrorCategory.FILE_IO,
                        ErrorSeverity.ERROR,
                        custom_message=error_msg,
                        custom_recovery=[
                            "Check available disk space",
                            "Verify file permissions",
                            "Try restarting the application"
                        ]
                    )
                else:
                    messagebox.showerror("Storage Error", error_msg)
                return None
            
            self.current_document = document
            self.logger.info(f"Successfully uploaded and processed document: {filename}")
            
            # Show success message briefly
            if progress:
                progress.update_message(f"Document '{filename}' processed successfully!")
                import time
                time.sleep(1)
            
            return document
            
        except MemoryError as e:
            error_msg = f"The document '{os.path.basename(file_path) if file_path else 'file'}' is too large to process."
            if self.error_handler:
                self.error_handler.handle_error(
                    e, ErrorCategory.MEMORY, ErrorSeverity.ERROR,
                    custom_message=error_msg,
                    custom_recovery=[
                        "Try a smaller document",
                        "Close other applications to free memory",
                        "Split large documents into smaller parts"
                    ]
                )
            else:
                messagebox.showerror("Memory Error", error_msg)
            return None
            
        except Exception as e:
            error_msg = f"An unexpected error occurred while processing the document."
            if self.error_handler:
                self.error_handler.handle_error(
                    e, ErrorCategory.DOCUMENT, ErrorSeverity.ERROR,
                    custom_message=error_msg,
                    context={"file_path": file_path, "operation": "upload_document"}
                )
            else:
                messagebox.showerror("Upload Error", f"{error_msg}: {str(e)}")
            self.logger.error(f"Error uploading document: {str(e)}")
            return None
            
        finally:
            if progress:
                progress.hide()
    
    @handle_errors(ErrorCategory.DOCUMENT, ErrorSeverity.ERROR, show_dialog=False, fallback_return="")
    def extract_text(self, file_path: str) -> str:
        """
        Extract text content from a document file with comprehensive error handling.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Extracted text content as string
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Document file not found: {file_path}")
        
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == '.pdf':
            return self._extract_pdf_text(file_path)
        elif file_extension in ['.doc', '.docx']:
            return self._extract_docx_text(file_path)
        elif file_extension == '.txt':
            return self._extract_txt_text(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}. Supported formats: {', '.join(self.SUPPORTED_FORMATS)}")
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF file using PyPDF2 with enhanced error handling."""
        text_content = ""
        failed_pages = []
        
        try:
            with open(file_path, 'rb') as file:
                try:
                    pdf_reader = PyPDF2.PdfReader(file)
                except PyPDF2.errors.PdfReadError as e:
                    if "password" in str(e).lower():
                        raise ValueError("PDF is password-protected. Please provide an unprotected version.")
                    else:
                        raise ValueError(f"PDF file is corrupted or invalid: {str(e)}")
                
                # Check if PDF has pages
                if len(pdf_reader.pages) == 0:
                    raise ValueError("PDF file contains no pages.")
                
                # Extract text from each page
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text_content += page_text + "\n"
                        else:
                            failed_pages.append(page_num + 1)
                    except Exception as e:
                        self.logger.warning(f"Could not extract text from page {page_num + 1}: {str(e)}")
                        failed_pages.append(page_num + 1)
                        continue
                
                # Warn about failed pages if any
                if failed_pages and self.error_handler:
                    if len(failed_pages) == len(pdf_reader.pages):
                        raise ValueError("Could not extract text from any pages. The PDF may contain only images or be corrupted.")
                    elif len(failed_pages) > len(pdf_reader.pages) * 0.5:
                        self.error_handler.handle_warning(
                            f"Could not extract text from {len(failed_pages)} out of {len(pdf_reader.pages)} pages. The document may contain images or be partially corrupted.",
                            ErrorCategory.DOCUMENT,
                            show_dialog=False
                        )
                        
        except FileNotFoundError:
            raise FileNotFoundError(f"PDF file not found: {file_path}")
        except PermissionError:
            raise PermissionError(f"Permission denied accessing PDF file: {file_path}")
        except Exception as e:
            if "password" in str(e).lower():
                raise ValueError("PDF is password-protected. Please provide an unprotected version.")
            else:
                raise ValueError(f"Error reading PDF file: {str(e)}")
        
        if not text_content.strip():
            raise ValueError("No readable text found in PDF. The document may contain only images.")
        
        return text_content.strip()
    
    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file using python-docx with enhanced error handling."""
        text_content = ""
        
        try:
            try:
                doc = DocxDocument(file_path)
            except Exception as e:
                if "not a zip file" in str(e).lower() or "bad zip file" in str(e).lower():
                    raise ValueError("Document file is corrupted or not a valid Word document.")
                elif "permission" in str(e).lower():
                    raise PermissionError(f"Permission denied accessing document: {file_path}")
                else:
                    raise ValueError(f"Cannot open Word document: {str(e)}")
            
            # Extract text from paragraphs
            paragraph_count = 0
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n"
                    paragraph_count += 1
            
            # Extract text from tables
            table_count = 0
            for table in doc.tables:
                table_text = ""
                for row in table.rows:
                    row_text = ""
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text += cell.text + " | "
                    if row_text:
                        table_text += row_text + "\n"
                if table_text:
                    text_content += table_text + "\n"
                    table_count += 1
            
            # Log extraction statistics
            self.logger.info(f"Extracted text from {paragraph_count} paragraphs and {table_count} tables")
                    
        except FileNotFoundError:
            raise FileNotFoundError(f"Word document not found: {file_path}")
        except PermissionError:
            raise PermissionError(f"Permission denied accessing Word document: {file_path}")
        except Exception as e:
            raise ValueError(f"Error reading Word document: {str(e)}")
        
        if not text_content.strip():
            raise ValueError("No readable text found in Word document. The document may be empty.")
        
        return text_content.strip()
    
    def _extract_txt_text(self, file_path: str) -> str:
        """Extract text from plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text_content = file.read()
            
            if not text_content.strip():
                raise ValueError("Text file is empty.")
            
            return text_content.strip()
            
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    text_content = file.read()
                return text_content.strip()
            except Exception as e:
                raise ValueError(f"Cannot read text file with any encoding: {str(e)}")
        except Exception as e:
            raise ValueError(f"Error reading text file: {str(e)}")
    
    def _validate_file(self, file_path: str) -> Tuple[bool, str]:
        """
        Validate uploaded file for format and size constraints.
        
        Args:
            file_path: Path to the file to validate
            
        Returns:
            Tuple of (is_valid, error_message)
        """
        try:
            # Check if file exists
            if not os.path.exists(file_path):
                return False, "File does not exist."
            
            # Check file extension
            file_extension = os.path.splitext(file_path)[1].lower()
            if file_extension not in self.SUPPORTED_FORMATS:
                return False, f"Unsupported file format. Supported formats: {', '.join(self.SUPPORTED_FORMATS)}"
            
            # Check file size
            file_size = os.path.getsize(file_path)
            if file_size > self.MAX_FILE_SIZE:
                return False, f"File too large. Maximum size: {self.MAX_FILE_SIZE // (1024*1024)}MB"
            
            if file_size == 0:
                return False, "File is empty."
            
            return True, ""
            
        except Exception as e:
            return False, f"Error validating file: {str(e)}"
    
    def get_current_document(self) -> Optional[Document]:
        """Get the currently loaded document."""
        return self.current_document
    
    def clear_current_document(self):
        """Clear the currently loaded document."""
        self.current_document = None
        self.logger.info("Cleared current document")
    
    def chunk_document(self, text: str, chunk_size: int = None, overlap: int = None) -> List[str]:
        """
        Split document text into overlapping chunks for efficient processing.
        
        Args:
            text: The text content to chunk
            chunk_size: Maximum characters per chunk (default: DEFAULT_CHUNK_SIZE)
            overlap: Characters to overlap between chunks (default: CHUNK_OVERLAP)
            
        Returns:
            List of text chunks
        """
        if not text.strip():
            return []
        
        chunk_size = chunk_size or self.DEFAULT_CHUNK_SIZE
        overlap = overlap or self.CHUNK_OVERLAP
        
        # Ensure overlap is not larger than chunk size
        overlap = min(overlap, chunk_size // 2)
        
        chunks = []
        start = 0
        
        while start < len(text):
            # Calculate end position
            end = start + chunk_size
            
            # If this is not the last chunk, try to break at a sentence or word boundary
            if end < len(text):
                # Look for sentence boundary (. ! ?) within the last 100 characters
                sentence_end = self._find_sentence_boundary(text, end - 100, end)
                if sentence_end > start:
                    end = sentence_end
                else:
                    # Look for word boundary within the last 50 characters
                    word_end = self._find_word_boundary(text, end - 50, end)
                    if word_end > start:
                        end = word_end
            
            # Extract chunk
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            # Move start position with overlap
            start = end - overlap
            
            # Prevent infinite loop
            if start >= end:
                start = end
        
        self.logger.info(f"Document chunked into {len(chunks)} chunks")
        return chunks
    
    def _find_sentence_boundary(self, text: str, start: int, end: int) -> int:
        """Find the last sentence boundary within the given range."""
        sentence_endings = re.finditer(r'[.!?]\s+', text[start:end])
        boundaries = list(sentence_endings)
        if boundaries:
            last_boundary = boundaries[-1]
            return start + last_boundary.end()
        return -1
    
    def _find_word_boundary(self, text: str, start: int, end: int) -> int:
        """Find the last word boundary within the given range."""
        word_boundaries = re.finditer(r'\s+', text[start:end])
        boundaries = list(word_boundaries)
        if boundaries:
            last_boundary = boundaries[-1]
            return start + last_boundary.start()
        return -1
    
    def _store_document(self, document: Document) -> bool:
        """
        Store a document in the persistent storage.
        
        Args:
            document: Document to store
            
        Returns:
            True if successful, False otherwise
        """
        try:
            # Check if document already exists (by filename and file_path)
            existing_doc = self.find_document_by_path(document.file_path)
            if existing_doc:
                # Update existing document
                self.stored_documents.remove(existing_doc)
            
            # Add new document
            self.stored_documents.append(document)
            
            # Save to persistent storage
            success = self.persistence_manager.save_documents(self.stored_documents)
            if success:
                self.logger.info(f"Document stored successfully: {document.filename}")
            else:
                self.logger.error(f"Failed to save document: {document.filename}")
            
            return success
            
        except Exception as e:
            self.logger.error(f"Error storing document: {str(e)}")
            return False
    
    def _load_stored_documents(self):
        """Load all stored documents from persistent storage."""
        try:
            self.stored_documents = self.persistence_manager.load_documents()
            self.logger.info(f"Loaded {len(self.stored_documents)} stored documents")
        except Exception as e:
            self.logger.error(f"Error loading stored documents: {str(e)}")
            self.stored_documents = []
    
    def get_stored_documents(self) -> List[Document]:
        """Get all stored documents."""
        return self.stored_documents.copy()
    
    def find_document_by_id(self, document_id: str) -> Optional[Document]:
        """Find a stored document by its ID."""
        for doc in self.stored_documents:
            if doc.id == document_id:
                return doc
        return None
    
    def find_document_by_path(self, file_path: str) -> Optional[Document]:
        """Find a stored document by its file path."""
        for doc in self.stored_documents:
            if doc.file_path == file_path:
                return doc
        return None
    
    def find_documents_by_filename(self, filename: str) -> List[Document]:
        """Find stored documents by filename (partial match)."""
        matching_docs = []
        filename_lower = filename.lower()
        for doc in self.stored_documents:
            if filename_lower in doc.filename.lower():
                matching_docs.append(doc)
        return matching_docs
    
    def remove_document(self, document_id: str) -> bool:
        """
        Remove a document from storage.
        
        Args:
            document_id: ID of the document to remove
            
        Returns:
            True if successful, False otherwise
        """
        try:
            doc_to_remove = self.find_document_by_id(document_id)
            if not doc_to_remove:
                return False
            
            self.stored_documents.remove(doc_to_remove)
            
            # Clear current document if it's the one being removed
            if self.current_document and self.current_document.id == document_id:
                self.current_document = None
            
            # Save updated list
            success = self.persistence_manager.save_documents(self.stored_documents)
            if success:
                self.logger.info(f"Document removed successfully: {doc_to_remove.filename}")
            
            return success
            
        except Exception as e:
            self.logger.error(f"Error removing document: {str(e)}")
            return False
    
    def find_relevant_chunks(self, query: str, max_chunks: int = 5) -> List[Tuple[str, str]]:
        """
        Find relevant document chunks based on a query using simple text matching.
        
        Args:
            query: Search query
            max_chunks: Maximum number of chunks to return
            
        Returns:
            List of tuples (chunk_text, document_filename)
        """
        if not self.current_document or not self.current_document.chunks:
            return []
        
        query_words = set(query.lower().split())
        chunk_scores = []
        
        for chunk in self.current_document.chunks:
            chunk_words = set(chunk.lower().split())
            # Simple scoring based on word overlap
            score = len(query_words.intersection(chunk_words))
            if score > 0:
                chunk_scores.append((score, chunk))
        
        # Sort by score and return top chunks
        chunk_scores.sort(key=lambda x: x[0], reverse=True)
        relevant_chunks = [(chunk, self.current_document.filename) 
                          for _, chunk in chunk_scores[:max_chunks]]
        
        return relevant_chunks
    
    def get_document_summary(self) -> str:
        """
        Get a summary of the current document.
        
        Returns:
            Summary string with document info
        """
        if not self.current_document:
            return "No document loaded."
        
        doc = self.current_document
        word_count = len(doc.text_content.split())
        char_count = len(doc.text_content)
        chunk_count = len(doc.chunks)
        
        return (f"Document: {doc.filename}\n"
                f"Size: {doc.file_size // 1024}KB\n"
                f"Words: {word_count:,}\n"
                f"Characters: {char_count:,}\n"
                f"Chunks: {chunk_count}\n"
                f"Uploaded: {doc.upload_date.strftime('%Y-%m-%d %H:%M:%S')}")
    
    def get_storage_summary(self) -> str:
        """
        Get a summary of all stored documents.
        
        Returns:
            Summary string with storage info
        """
        if not self.stored_documents:
            return "No documents stored."
        
        total_size = sum(doc.file_size for doc in self.stored_documents)
        total_chunks = sum(len(doc.chunks) for doc in self.stored_documents)
        
        summary = f"Stored Documents: {len(self.stored_documents)}\n"
        summary += f"Total Size: {total_size // 1024}KB\n"
        summary += f"Total Chunks: {total_chunks}\n\n"
        
        for doc in self.stored_documents:
            summary += f"• {doc.filename} ({len(doc.chunks)} chunks)\n"
        
        return summary  
  
    def _handle_memory_optimization(self, optimization_type: str):
        """
        Handle memory optimization requests
        
        Args:
            optimization_type: Type of optimization requested
        """
        try:
            if optimization_type in ["low_memory_optimization", "comprehensive_optimization"]:
                self._optimize_document_memory()
                
        except Exception as e:
            self.logger.error(f"Error in document memory optimization: {e}")
    
    def _optimize_document_memory(self):
        """Optimize document memory usage"""
        self.logger.info("Optimizing document processor memory usage")
        
        # Use memory optimizer to optimize stored documents
        optimized_count = self.memory_optimizer.optimize_document_memory(
            self.stored_documents, 
            force=True
        )
        
        if optimized_count > 0:
            self.logger.info(f"Optimized {optimized_count} documents for memory usage")
        
        # Clear current document content if memory pressure is high
        if self.memory_optimizer.check_memory_pressure() and self.current_document:
            if hasattr(self.current_document, 'text_content'):
                # Keep chunks but clear full text content
                self.current_document.text_content = None
                self.logger.info("Cleared current document text content to save memory")
    
    def upload_document_async(self, callback: Optional[Callable] = None, 
                             error_callback: Optional[Callable] = None,
                             progress_callback: Optional[Callable] = None) -> str:
        """
        Upload document asynchronously to maintain UI responsiveness
        
        Args:
            callback: Success callback function
            error_callback: Error callback function
            progress_callback: Progress callback function
            
        Returns:
            Task ID for tracking
        """
        def upload_task(progress_callback=None):
            """Background task for document upload"""
            if progress_callback:
                progress_callback(0.1, "Opening file dialog...")
            
            # This would need to be adapted for async file dialog
            # For now, we'll use the synchronous version
            return self.upload_document()
        
        return self.background_processor.submit_document_processing_task(
            document_path="user_selected",  # Placeholder
            processor_function=upload_task,
            callback=callback,
            error_callback=error_callback,
            progress_callback=progress_callback
        )
    
    def extract_text_async(self, file_path: str, 
                          callback: Optional[Callable] = None,
                          error_callback: Optional[Callable] = None,
                          progress_callback: Optional[Callable] = None) -> str:
        """
        Extract text from document asynchronously
        
        Args:
            file_path: Path to document file
            callback: Success callback function
            error_callback: Error callback function
            progress_callback: Progress callback function
            
        Returns:
            Task ID for tracking
        """
        def extract_task(file_path, progress_callback=None):
            """Background task for text extraction"""
            if progress_callback:
                progress_callback(0.2, "Starting text extraction...")
            
            # Record start time for performance monitoring
            start_time = time.time()
            
            try:
                text_content = self.extract_text(file_path)
                
                # Record processing time
                processing_time = (time.time() - start_time) * 1000  # Convert to ms
                self.performance_monitor.record_response_time(processing_time)
                
                if progress_callback:
                    progress_callback(1.0, "Text extraction completed")
                
                return text_content
                
            except Exception as e:
                self.logger.error(f"Error in async text extraction: {e}")
                raise
        
        return self.background_processor.submit_document_processing_task(
            document_path=file_path,
            processor_function=extract_task,
            callback=callback,
            error_callback=error_callback,
            progress_callback=progress_callback
        )
    
    def chunk_document_optimized(self, text: str, chunk_size: int = None, 
                                overlap: int = None, 
                                progress_callback: Optional[Callable] = None) -> List[str]:
        """
        Optimized document chunking with memory management
        
        Args:
            text: Text content to chunk
            chunk_size: Maximum characters per chunk
            overlap: Characters to overlap between chunks
            progress_callback: Progress callback function
            
        Returns:
            List of text chunks
        """
        if not text.strip():
            return []
        
        chunk_size = chunk_size or self.DEFAULT_CHUNK_SIZE
        overlap = overlap or self.CHUNK_OVERLAP
        
        # Use memory optimizer for string optimization
        optimized_text = self.memory_optimizer.optimize_string_memory([text])[0]
        
        # Process in batches to manage memory
        chunks = []
        start = 0
        total_length = len(optimized_text)
        processed_chars = 0
        
        while start < total_length:
            # Calculate batch end
            batch_end = min(start + chunk_size * self.chunk_processing_batch_size, total_length)
            batch_text = optimized_text[start:batch_end]
            
            # Process batch
            batch_chunks = self._chunk_text_batch(batch_text, chunk_size, overlap, start)
            chunks.extend(batch_chunks)
            
            # Update progress
            processed_chars += len(batch_text)
            if progress_callback:
                progress = processed_chars / total_length
                progress_callback(progress, f"Processed {processed_chars}/{total_length} characters")
            
            # Move to next batch
            start = batch_end - overlap
            
            # Check memory pressure and optimize if needed
            if self.memory_optimizer.check_memory_pressure():
                self.memory_optimizer.force_garbage_collection()
        
        self.logger.info(f"Document chunked into {len(chunks)} chunks (optimized)")
        return chunks
    
    def _chunk_text_batch(self, text: str, chunk_size: int, overlap: int, offset: int = 0) -> List[str]:
        """
        Chunk a batch of text
        
        Args:
            text: Text to chunk
            chunk_size: Chunk size
            overlap: Overlap size
            offset: Offset for boundary detection
            
        Returns:
            List of chunks
        """
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            
            # Find good boundary if not at end
            if end < len(text):
                # Look for sentence boundary
                sentence_end = self._find_sentence_boundary(text, max(0, end - 100), end)
                if sentence_end > start:
                    end = sentence_end
                else:
                    # Look for word boundary
                    word_end = self._find_word_boundary(text, max(0, end - 50), end)
                    if word_end > start:
                        end = word_end
            
            # Extract chunk
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            # Move to next chunk with overlap
            start = end - overlap
            if start >= end:
                start = end
        
        return chunks
    
    def get_processing_statistics(self) -> Dict[str, Any]:
        """
        Get document processing statistics
        
        Returns:
            Dictionary with processing statistics
        """
        try:
            memory_usage = self.memory_optimizer.get_current_memory_usage()
            
            # Calculate document memory usage
            doc_memory_mb = 0.0
            for doc in self.stored_documents:
                if hasattr(doc, 'text_content') and doc.text_content:
                    doc_memory_mb += len(doc.text_content.encode('utf-8')) / (1024 * 1024)
                if hasattr(doc, 'chunks') and doc.chunks:
                    for chunk in doc.chunks:
                        doc_memory_mb += len(chunk.encode('utf-8')) / (1024 * 1024)
            
            return {
                'stored_documents': len(self.stored_documents),
                'current_document_loaded': self.current_document is not None,
                'total_memory_usage_mb': memory_usage,
                'document_memory_mb': round(doc_memory_mb, 2),
                'memory_pressure': self.memory_optimizer.check_memory_pressure(),
                'max_documents_in_memory': self.max_documents_in_memory,
                'chunk_batch_size': self.chunk_processing_batch_size
            }
            
        except Exception as e:
            self.logger.error(f"Error getting processing statistics: {e}")
            return {'error': str(e)}
    
    def optimize_storage(self):
        """Optimize document storage for memory efficiency"""
        self.logger.info("Optimizing document storage")
        
        # Remove excess documents from memory
        while len(self.stored_documents) > self.max_documents_in_memory:
            # Remove oldest document (simple strategy)
            oldest_doc = min(
                self.stored_documents,
                key=lambda d: getattr(d, 'upload_date', datetime.min)
            )
            
            # Clear content but keep metadata
            if hasattr(oldest_doc, 'text_content'):
                oldest_doc.text_content = None
            if hasattr(oldest_doc, 'chunks'):
                oldest_doc.chunks = []
            
            self.logger.info(f"Optimized storage for document: {oldest_doc.filename}")
        
        # Force garbage collection
        self.memory_optimizer.force_garbage_collection()
    
    def set_memory_limits(self, max_documents: int, batch_size: int):
        """
        Set memory limits for document processing
        
        Args:
            max_documents: Maximum documents to keep in memory
            batch_size: Batch size for chunk processing
        """
        self.max_documents_in_memory = max_documents
        self.chunk_processing_batch_size = batch_size
        
        self.logger.info(f"Document processor limits updated: max_docs={max_documents}, batch_size={batch_size}")
        
        # Optimize current storage if needed
        if len(self.stored_documents) > max_documents:
            self.optimize_storage()